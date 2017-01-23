import docx
from src.user_input import last_name, husband, spouse
from src.date import current_date, save_date
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_BREAK

def document():
    new_document = docx.Document() # Creates document to memory


    # def for spouse
    def spouse_attend():
        if spouse == "None":
            pass
        else:
            attendees_three = new_document.add_paragraph(text='{} {} (Client - Spouse)'.format(spouse, last_name), style=None)

    def spouse_attend_two():
        if spouse == "None":
            pass
        else:
            # Second Client Information
            client_2 = new_document.add_paragraph(text='{} {}:'.format(spouse, last_name), style=None)  # Client two information
            client_2.runs[0].bold = True
            client_2_info = new_document.add_paragraph(text='Add Info Here!\n')
            client_2_info.style = 'List Bullet'
            client_2_info.paragraph_format.left_indent = Inches(.75)


    # Below is the heading of the document
    heading = new_document.add_paragraph(text="{} Meeting Minutes\n".format(last_name), style=None) # Heading of Document
    heading.new_document = WD_LINE_SPACING.DOUBLE
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(16)
    heading.add_run(current_date)
    heading.runs[1].bold = True
    heading.add_run('\nStart: (Add Time!) End: (Add Time!)')
    heading.add_break = WD_BREAK.TEXT_WRAPPING
    heading.add_run('\nTotal Meeting Time: (Add Total Time)\n')
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



    # The attendees at the meeting
    attendees = new_document.add_paragraph(text="Attendees:", style=None) # Attendees section of the document
    attendees.runs[0].bold = True

    attendees_one = new_document.add_paragraph(text='Jerry Graham (Adviser)', style=None)

    attendees_two = new_document.add_paragraph(text='Justin Kendter (Assistant)', style=None)

    attendees_three = new_document.add_paragraph(text='{} {} (Client - Husband)'.format(husband, last_name), style=None)

    spouse_attend() # Grabs def above for spouse



    # This is Client information
    client_information = new_document.add_paragraph(text='\nClient Information:', style=None) # Information on the client
    client_information.runs[0].bold = True
    client_information_info = new_document.add_paragraph(text='Add Info Here!\n')
    client_information_info.style = 'List Bullet'
    client_information_info.paragraph_format.left_indent = Inches(.75)


    # First Client Information
    client_1 = new_document.add_paragraph(text='Justin Kendter:', style=None) # Client one information
    client_1.runs[0].bold = True
    client_1_info = new_document.add_paragraph(text='Add Info Here!\n')
    client_1_info.style = 'List Bullet'
    client_1_info.paragraph_format.left_indent = Inches(.75)

    spouse_attend_two() # Calls for 2nd spouse


    # Pros and Cons Table
    # Client Pros
    concerns_pros = new_document.add_table(rows=1, cols=2)
    pro_con_cells = concerns_pros.rows[0].cells
    pro_con_cells[0].paragraphs[0].add_run('Client Concerns:').bold = True
    # Client Cons
    pro_con_cells[1].paragraphs[0].add_run('Client Pros:').bold=True



    concerns = pro_con_cells[0].add_paragraph(text='Add Info Here!', style=None)
    concerns.style = 'List Bullet'
    concerns.paragraph_format.left_indent = Inches(.75)

    pros = pro_con_cells[1].add_paragraph(text='Add Info Here!', style=None)
    pros.style = 'List Bullet'
    pros.paragraph_format.left_indent = Inches(.75)



    # Advisors Advice Section
    adviser_advice = new_document.add_paragraph(text='\nAdvisers Advice:', style=None)
    adviser_advice.runs[0].bold = True
    adviser_advice_info = new_document.add_paragraph(text='Add Info Here!\n')
    adviser_advice_info.style = 'List Bullet'
    adviser_advice_info.paragraph_format.left_indent = Inches(.75)



    # Advisers action items.
    adviser_action_items = new_document.add_paragraph(text='Adviser Action Items:', style=None)
    adviser_action_items.runs[0].bold = True
    adviser_action_items_info = new_document.add_paragraph(text='Add Info Here!\n')
    adviser_action_items_info.style = 'List Bullet'
    adviser_action_items_info.paragraph_format.left_indent = Inches(.75)



    #Client financial information
    current_client_financials = new_document.add_paragraph(text='Current Client Financial\'s:', style =None)
    current_client_financials.runs[0].bold = True
    current_client_financials_info = new_document.add_paragraph(text='Add Info Here!\n')
    current_client_financials_info.style = 'List Bullet'
    current_client_financials_info.paragraph_format.left_indent = Inches(.75)



    try:
        save_name = '{} Meeting Minutes {}.docx'.format(last_name, save_date)
        save_path = 'C:\\Users\\Justin\\Documents\\Client Meeting Minutes\\'
        new_document.save(save_path + save_name)
        save_location = save_path + save_name

        def getText(filename):
            doc = docx.Document(filename)
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            return '\n'.join(fullText)

        print(getText(save_location))
        print('\nSaved To:', save_location)
    except:
        print('Document is open')